from django.shortcuts import render, redirect
from django.http import HttpResponseBadRequest, FileResponse
from django.core.paginator import Paginator
from docx import Document
from html2text import html2text
from io import BytesIO
import markdown

from .forms import WrittingForm
from .get_result import get_writing_result
from .models import UserWrittings
from user_auth.models import UserProfile
from django.contrib.auth.models import User

def home_page(request):
    form = WrittingForm()
    if request.user.is_authenticated and not hasattr(request.user, 'userprofile'):
        UserProfile.objects.create(user=request.user)
    if request.method == 'POST':
        form = WrittingForm(data=request.POST)
        if form.is_valid():
            task = request.POST.get('task')
            writting = request.POST.get('writting')
            # Store task and writting in the session
            request.session['task'] = task
            request.session['writting'] = writting
            request.session['new'] = True

            return redirect('writting-result')
        
    sample_writtings = UserWrittings.objects.filter(public_status=True).order_by('?')[:9]
        
    return render(request, 'home_page.html', {'form': form, 'sample_writtings': sample_writtings})

def result(request):
    if request.method == 'POST':
        writting_id = request.GET.get('writting_id')
        request.session['new'] = False
        return redirect(f'/result?writting_id={writting_id}')
    else:
        new_status = request.session.get('new')
        if new_status:
            task = request.session.get('task')
            writting = request.session.get('writting')
            result = get_writing_result(task, html2text(writting))
            result['feed_back'] = markdown.markdown(result['feed_back']).replace('\n', '<br>')

            writting_data = UserWrittings(
                user_name=request.user,
                task=task,
                writting=writting,
                score=result['score'],
                feed_back=result['feed_back'],
                public_status=False 
            )
            writting_data.save()
        else:
            writting_id = request.GET.get('writting_id')
            writting_data = UserWrittings.objects.get(id=writting_id)

        form = WrittingForm(data={'writting': writting_data.writting})
        context = {'writting_data': writting_data, 'form': form}
        return render(request, 'writting_result.html', context)

def public_writting(request):
    if request.method == 'POST':
        task = request.POST.get('task')
        writting = request.POST.get('writting')
        # Store task and writting in the session
        request.session['task'] = task
        request.session['writting'] = writting
        request.session['new'] = False

        return redirect('writting-result')
    else:
        # public_writtings = UserWrittings.objects.filter(public_status=True)
        p = Paginator(UserWrittings.objects.filter(public_status=True), 6)
        page = request.GET.get('page')
        public_writtings = p.get_page(page)

        print(public_writtings)
        return render(request, 'public_writting.html', {'public_writtings': public_writtings})

# @login_required(login_url='user-login')
def writting_history(request):

    user_id = int(request.GET.get('user_id'))
    user = User.objects.get(id=user_id)

    if request.user.id == user_id:
        user_data = UserWrittings.objects.filter(user_name=user)
    else:
        user_data = UserWrittings.objects.filter(user_name=user, public_status=True)

    return render(request, 'writting_history.html', {'user_data': user_data, 'user': user})

def toggle_public(request):
    if request.method == 'POST':
        writting_id = request.POST.get('id')
        writting = UserWrittings.objects.get(id=writting_id)
        writting.public_status = not writting.public_status
        writting.save()
        return redirect(f"/history?user_id={writting.user_name.id}")
    else:
        return HttpResponseBadRequest("This URL only supports POST requests")

def download_doc(request):
    if request.method == 'POST':
        doc = Document()
        doc.add_heading('Task', 0)
        doc.add_paragraph(request.POST.get('task'))

        doc.add_heading('Writting', 0)
        doc.add_paragraph(html2text(request.POST.get('writting')))

        doc.add_heading('Result', level=0)
        doc.add_heading('Score', 1)
        doc.add_paragraph(request.POST.get('score'))

        doc.add_heading('Feedback', 1)
        doc.add_paragraph(html2text(request.POST.get('feed_back')))

        # Save the Document to a BytesIO object
        f = BytesIO()
        doc.save(f)
        f.seek(0)

        # Create a FileResponse with the BytesIO object and set the Content-Disposition header
        response = FileResponse(f, as_attachment=True, filename='writting_feedback.docx')
        return response
    else:
        return HttpResponseBadRequest("This URL only supports POST requests")